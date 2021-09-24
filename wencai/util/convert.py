import os
import re
from wencai.util.utils import ech, sort_filename
from PIL import Image
from tqdm import tqdm


def create_dest(dest, type=None):
    if not dest:
        dest = os.path.join(os.getcwd(), f'combined{type}')
    os.makedirs(os.path.dirname(dest), exist_ok=True)
    return dest


def pdf2jpg(pdfPath, imagePath, zoom=1):
    import fitz
    ech(f"pdf2jpg: {pdfPath} -> {imagePath}", 'green')
    pdfDoc = fitz.open(pdfPath)
    for pg in tqdm(range(pdfDoc.pageCount)):
        page = pdfDoc[pg]
        rotate = int(0)
        # 每个尺寸的缩放系数为1.3，这将为我们生成分辨率提高2.6的图像。
        # 此处若是不做设置，默认图片大小为：792X612, dpi=96    默认 600*800
        mat = fitz.Matrix(zoom * 2, zoom * 2).preRotate(rotate)
        pix = page.getPixmap(matrix=mat, alpha=False)

        if not os.path.exists(imagePath):  # 判断存放图片的文件夹是否存在
            os.makedirs(imagePath)  # 若图片文件夹不存在就创建

        # 将图片写入指定的文件夹内
        pix.writePNG(os.path.join(imagePath, f'images_{pg}.png'))


def jpg2pdf(img_path, dest=None):
    """将路径下的所有图片转换成一个pdf，目的为None时转换到当前文件夹下"""
    import fitz
    dest = create_dest(dest, '.pdf')
    ech(f'jpg2pdf: {img_path} -> {dest}', 'green')
    doc = fitz.open()
    imgs = [i for i in os.listdir(img_path) if i[-4:] in ['.jpg', '.png']]

    for img in sort_filename(imgs):
        img_file = os.path.join(img_path, img)
        imgdoc = fitz.open(img_file)
        pdfbytes = imgdoc.convertToPDF()
        imgpdf = fitz.open(img_file[:-4] + '.pdf', pdfbytes)
        doc.insertPDF(imgpdf)
    doc.save(dest)
    doc.close()


def pdf2txt(pdf_path):
    from pdfminer.pdfparser import PDFParser
    from pdfminer.pdfdocument import PDFDocument
    from pdfminer.pdfpage import PDFPage
    from pdfminer.pdfpage import PDFTextExtractionNotAllowed
    from pdfminer.pdfinterp import PDFResourceManager
    from pdfminer.pdfinterp import PDFPageInterpreter
    from pdfminer.layout import LAParams, LTTextBoxHorizontal
    from pdfminer.converter import PDFPageAggregator
    fp = open(pdf_path, 'rb')
    s = ''
    # 来创建一个pdf文档分析器
    parser = PDFParser(fp)
    # 创建一个PDF文档对象存储文档结构
    document = PDFDocument(parser)
    # 检查文件是否允许文本提取
    if not document.is_extractable:
        raise PDFTextExtractionNotAllowed
    else:
        # 创建一个PDF资源管理器对象来存储共赏资源
        rsrcmgr = PDFResourceManager()
        # 设定参数进行分析
        laparams = LAParams()
        # 创建一个PDF设备对象
        # device=PDFDevice(rsrcmgr)
        device = PDFPageAggregator(rsrcmgr, laparams=laparams)
        # 创建一个PDF解释器对象
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        # 处理每一页
        for page in PDFPage.create_pages(document):
            interpreter.process_page(page)
            # 接受该页面的LTPage对象
            layout = device.get_result()
            for x in layout:
                if(isinstance(x, LTTextBoxHorizontal)):
                    s += x.get_text()
    return s


def html2txt(filename):
    from bs4 import BeautifulSoup
    f = open(filename, "rb")
    return BeautifulSoup(f).get_text()


def word2txt(filename):
    from docx import Document
    document = Document(filename)
    return '\n'.join([paragraph.text for paragraph in document.paragraphs])


def embedded_numbers(s):
    re_digits = re.compile(r'(\d+)')
    pieces = re_digits.split(s)                 # 切成数字和非数字
    pieces[1::2] = map(int, pieces[1::2])       # 将数字部分转成整数
    return pieces


def sort_string(lst):
    return sorted(lst, key=embedded_numbers)    # 将前面的函数作为key来排序


def get_file_list(src, recursive):
    all_path = []
    if recursive:
        # root文件夹的路径  dirs 路径下的文件夹列表  files路径下的文件列表
        for root, dirs, files in os.walk(src):
            for f in files:
                if f[-4:] in ['.jpg', '.png']:  # 子串在母串里面不
                    all_path.append(os.path.join(root, f))
    else:
        for f in os.listdir(src):
            if f[-4:] in ['.jpg', '.png']:  # 子串在母串里面不
                all_path.append(os.path.join(src, f))
    return sort_filename(all_path)


def get_size(pic, vertical):
    """返回图片（横向，纵向）"""
    a, b = Image.open(pic).size
    if vertical:
        return min(a, b), max(a, b)
    return max(a, b), min(a, b)


def concat_pic_pdf(src, dest=None, line_max=2, recursive=False, alpha=1, vertical=False, row_max=None):
    """将src文件夹下图片合成到dest（pdf文件），每行line_max个，扩大系数为alpha，vertical代表为竖直状态"""
    num = 0
    dest = create_dest(dest, '.pdf')
    concat_tmp = os.path.join(src, 'concat')
    import shutil
    shutil.rmtree(concat_tmp, ignore_errors=True)
    os.makedirs(concat_tmp)

    ech(f'concat_pic: {src} -> {concat_tmp} -> {dest}', 'green')
    ech("src paths are:", 'yellow')
    all_path = get_file_list(src, recursive)
    print(all_path)
    N = len(all_path)
    assert N > 0

    # 固定一个宽度，最后的实际宽度（取均值）
    width = 0
    for i in range(N):
        w, _ = get_size(all_path[i], vertical)
        width += w
    width = int(width * alpha/N)

    # 同行的高度以第一个为准
    for i in tqdm(range(0, N, line_max)):
        w, h = get_size(all_path[i], vertical)
        h = int(width * h/w)
        # print(f"the {i}th height is {h}")

        toImage = Image.new('RGBA', (width * line_max, h))

        for j in range(line_max):
            # 每次打开图片绝对路路径列表的第一张图片
            pic_fole_head = Image.open(all_path[num])
            w, h = get_size(all_path[num], vertical)

            # 按照指定的尺寸，给图片重新赋值，<PIL.Image.Image image mode=RGB size=200x200 at 0x127B7978>
            tmppic = pic_fole_head.resize((width, h))

            # 计算每个图片的左上角的坐标点(0, 0)，(0, 200)，(0, 400)，(200, 0)，(200, 200)。。。。(400, 400)
            loc = (int(j * width), 0)
            toImage.paste(tmppic, loc)
            num = num + 1
            if num >= N:
                break

        toImage.save(f'{concat_tmp}/{i}.png')

        row = int(i/line_max)
        if row_max and (row + 1) % row_max == 0:
            jpg2pdf(concat_tmp, dest.replace('.pdf', f'-{row}.pdf'))
            shutil.rmtree(concat_tmp, ignore_errors=True)
            os.makedirs(concat_tmp)

    if os.listdir(concat_tmp):
        jpg2pdf(concat_tmp, dest)


def concat_pic(src, dest=None, line_max=2, recursive=False, alpha=1, vertical=False):
    """将src文件夹下图片合成到dest（pdf文件），每行line_max个，扩大系数为alpha，vertical代表为竖直状态"""
    num = 0
    dest = create_dest(dest, '.png')
    ech(f'concat_pic: {src} -> {dest}', 'green')

    ech("src paths are:", 'yellow')
    all_path = get_file_list(src, recursive)
    print(all_path)
    N = len(all_path)
    assert N > 0
    row_max = int((N-1) / line_max) + 1

    # 固定一个宽度，最后的实际宽度（取均值）
    width = 0
    for i in range(N):
        w, _ = get_size(all_path[i], vertical)
        width += w
    width = int(width * alpha/N)

    # 同行的高度以第一个为准
    height_total = 0
    for i in range(0, N, line_max):
        w, h = get_size(all_path[i], vertical)
        h = int(width * h/w)
        # print(f"the {i}th height is {h}")
        height_total += h

    toImage = Image.new('RGBA', (width * line_max, height_total))
    height = 0
    height_row = 0

    i = -1
    while True:
        i += 1
        for j in range(line_max):
            # 每次打开图片绝对路路径列表的第一张图片
            pic_fole_head = Image.open(all_path[num])
            w, h = get_size(all_path[num], vertical)

            # 获取行首图片的高度
            if j % line_max == 0:
                height += height_row
                height_row = int(width * h/w)

            # 按照指定的尺寸，给图片重新赋值，<PIL.Image.Image image mode=RGB size=200x200 at 0x127B7978>
            tmppic = pic_fole_head.resize((width, height_row))

            # 计算每个图片的左上角的坐标点(0, 0)，(0, 200)，(0, 400)，(200, 0)，(200, 200)。。。。(400, 400)
            loc = (int(j * width), int(height))
            print(f"the {num+1}th/{N} place at: {str(loc)}")
            toImage.paste(tmppic, loc)
            num = num + 1

            if num >= N:
                break
        if num >= N:
            break

    ech(f'concat image size is: {toImage.size}', 'yellow')
    toImage.save(dest)


if __name__ == "__main__":
    # pdf2jpg('/home/sunie/Desktop/material/convert/core_qmof_pdf/zaac.201500654.pdf', '/tmp/test')
    concat_pic_pdf('/tmp/test', '/tmp/test/concat/img.pdf', vertical=True)
    # jpg2pdf('/tmp/test/concat', '/tmp/test/concat/merged.pdf')
